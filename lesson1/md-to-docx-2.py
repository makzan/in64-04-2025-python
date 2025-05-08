# This program reads all the .md files in the md folder 
# and converts them to docx files in the docx folder.
# It reads each .md line by line, convert headings to headings
# and normal lines to paragraphs.
# Use os, glob, docx. 

import os
import glob
from docx import Document

source_folder = 'md'
source_extention = '.md'
target_folder = 'docx'

# Read all .md files in the md folder
md_files = glob.glob(f'{source_folder}/*{source_extention}')

# Convert each .md file to .docx
for md_file in md_files:
    print(f"Converting {md_file} to .docx")
    # Read the .md file
    with open(md_file, 'r') as f:
        lines = f.readlines()

    # Create a new .docx file
    doc = Document()

    # Convert headings to headings and normal lines to paragraphs
    for line in lines:
        if line.strip() == '':
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)

    # Save the .docx file
    
    filename = os.path.basename(md_file).replace('.md', '.docx')
    doc.save(f'{target_folder}/{filename}')

    print(f"Done converting {md_file} to .docx")